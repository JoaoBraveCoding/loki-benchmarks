import os
import sys
import argparse
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
from bs4 import BeautifulSoup

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def add_table_of_contents(soup, doc):
    toc = soup.find('ul')
    if toc:
        for li in toc.find_all('li'):
            link = li.find('a')
            if link and link['href'].startswith('#'):
                heading_text = link.text
                toc_paragraph = doc.add_paragraph()
                add_hyperlink(toc_paragraph, f'#{heading_text}', heading_text)

def add_markdown_to_docx(md_content, doc, base_path):
    html = markdown.markdown(md_content)
    soup = BeautifulSoup(html, 'html.parser')

    heading_map = {}
    toc_inserted = False

    for element in soup:
        if element.name == 'h1':
            paragraph = doc.add_heading(element.text, level=1)
            heading_map[element.text] = paragraph
        elif element.name == 'h2':
            paragraph = doc.add_heading(element.text, level=2)
            heading_map[element.text] = paragraph
            if element.text.lower() == 'table of contents' and not toc_inserted:
                add_table_of_contents(soup, doc)
                toc_inserted = True
        elif element.name == 'h3':
            paragraph = doc.add_heading(element.text, level=3)
            heading_map[element.text] = paragraph
        elif element.name == 'p':
            paragraph = doc.add_paragraph(element.text)
            for img in element.find_all('img'):
                img_src = img['src'].lstrip('./')
                img_path = os.path.join(base_path, img_src)
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(5.0))
                else:
                    paragraph.add_run(f"[Image not found: {img_path}]")
        elif element.name == 'ul' and not toc_inserted:
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListBullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListNumber')
        elif element.name == 'a':
            paragraph = doc.add_paragraph()
            add_hyperlink(paragraph, element['href'], element.text)

    for heading_text, paragraph in heading_map.items():
        bookmark = OxmlElement('w:bookmarkStart')
        bookmark.set(qn('w:id'), str(hash(heading_text)))
        bookmark.set(qn('w:name'), heading_text)
        paragraph._p.insert(0, bookmark)
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(hash(heading_text)))
        paragraph._p.append(bookmark_end)

def convert_readme_to_docx(readme_dir, output_path):
    readme_path = os.path.join(readme_dir, 'README.md')
    if not os.path.exists(readme_path):
        print(f"README.md not found in {readme_dir}")
        return

    with open(readme_path, 'r') as file:
        md_content = file.read()

    doc = Document()
    add_markdown_to_docx(md_content, doc, readme_dir)
    doc.save(output_path)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Convert a README.md file to a DOCX file.')
    parser.add_argument('readme_dir', type=str, help='Directory containing the README.md file')
    args = parser.parse_args()

    readme_dir = args.readme_dir
    output_path = os.path.join(readme_dir, 'README.docx')
    convert_readme_to_docx(readme_dir, output_path)
    print(f"Converted README.md in {readme_dir} to {output_path}")